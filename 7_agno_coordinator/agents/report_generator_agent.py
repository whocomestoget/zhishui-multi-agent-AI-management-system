#!/usr/bin/env python3
"""
智水信息Multi-Agent智能分析系统 - Report Generator Agent
报告生成专家智能体

功能职责：
1. 聚合多个Agent的分析结果
2. 基于模板生成HTML可视化报告
3. 提供综合性的决策建议
4. 生成交互式数据可视化

Author: 商海星辰队
Version: 1.0.0
"""

import json
import logging
from datetime import datetime
from typing import Dict, List, Any, Optional
from dataclasses import dataclass
from openai import OpenAI
from .base_agent import BaseAgent, AgentTask, AgentResult
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

# ================================
# 1. 配置和初始化
# ================================

# AI模型配置
AI_CONFIG = {
    "api_key": "sk-Wy5BpzceSjET0ZiZWvaMaxUTrUiEKYGgElx10VL88lAnhgSe",
    "api_base": "http://38.246.251.165:3002/v1",
    "model": "gemini-2.5-pro",
    "temperature": 0.7,
    "max_tokens": 65000,
}

# 初始化OpenAI客户端
client = OpenAI(
    api_key=AI_CONFIG["api_key"],
    base_url=AI_CONFIG["api_base"]
)

# 日志配置
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("ReportGeneratorAgent")

# ================================
# 2. 数据结构定义
# ================================

@dataclass
class AgentAnalysisResult:
    """单个Agent分析结果"""
    agent_name: str
    agent_type: str
    analysis_data: Dict[str, Any]
    confidence_score: float
    execution_time: float
    timestamp: str
    recommendations: List[str]
    key_insights: List[str]

@dataclass
class ReportSection:
    """报告章节结构"""
    section_id: str
    title: str
    content: str
    charts: List[Dict[str, Any]]
    priority: int
    data_sources: List[str]

@dataclass
class ComprehensiveReport:
    """综合分析报告"""
    report_id: str
    title: str
    executive_summary: str
    sections: List[ReportSection]
    overall_recommendations: List[str]
    risk_assessment: Dict[str, Any]
    next_actions: List[str]
    confidence_metrics: Dict[str, float]
    generation_timestamp: str

# ================================
# 3. Report Generator Agent实现
# ================================

class ReportGeneratorAgent(BaseAgent):
    """报告生成专家智能体
    
    专业职责：
    - 聚合多Agent分析结果
    - 生成HTML可视化报告
    - 提供综合决策建议
    - 创建交互式图表
    """
    
    def __init__(self):
        super().__init__("report_generator", "报告生成专家")
        self.agent_type = "report_generator"
        self.version = "1.0.0"
        
        # 报告模板配置
        self.report_templates = self._load_report_templates()
        self.chart_configs = self._load_chart_configurations()
        
        logger.info(f"初始化 {self.agent_name} 完成")
    
    def get_required_fields(self) -> List[str]:
        """获取必需的输入字段"""
        return ["agent_results", "output_mode"]
    
    def validate_input_data(self, task: AgentTask) -> tuple[bool, List[str]]:
        """验证输入数据"""
        errors = []
        task_data = task.input_data
        
        # 兼容处理：从input_data字段或直接从task_data获取业务数据
        business_data = task_data.get("input_data", task_data)
        
        if "agent_results" not in business_data:
            errors.append("缺少agent_results字段")
        elif not isinstance(business_data["agent_results"], list):
            errors.append("agent_results必须是列表类型")
        
        if "output_mode" not in business_data:
            errors.append("缺少output_mode字段")
        elif business_data["output_mode"] not in ["word", "html", "dialog", "all"]:
            errors.append("output_mode必须是word、html、dialog或all之一")
        
        return len(errors) == 0, errors
    
    def perform_analysis(self, data: Dict[str, Any], task: AgentTask) -> Dict[str, Any]:
        """执行分析"""
        agent_results = self._parse_agent_results(data.get("agent_results", []))
        output_mode = data.get("output_mode", "word")
        
        if output_mode == "all":
            return self.generate_all_modes_analysis(agent_results)
        else:
            comprehensive_analysis = self._generate_comprehensive_analysis(agent_results, output_mode)
            
            result = {
                "comprehensive_analysis": comprehensive_analysis,
                "agent_results": agent_results,
                "output_mode": output_mode
            }
            
            if output_mode == "word":
                result["word_file_path"] = self._generate_word_report(comprehensive_analysis, agent_results)
            elif output_mode == "html":
                result["html_content"] = self._generate_html_report(comprehensive_analysis, agent_results)
            elif output_mode == "dialog":
                result["dialog_response"] = comprehensive_analysis.get("executive_summary", "")
            
            return result
    
    def calculate_confidence_score(self, result: Dict[str, Any]) -> float:
        """计算置信度分数"""
        if "comprehensive_analysis" in result:
            return result["comprehensive_analysis"].get("overall_confidence", 0.8)
        return 0.8
    
    def generate_recommendations(self, result: Dict[str, Any]) -> List[str]:
        """生成建议"""
        if "comprehensive_analysis" in result:
            return result["comprehensive_analysis"].get("recommendations", [])
        return []
    
    def get_system_prompt(self) -> str:
        """获取系统提示词"""
        return """
你是智水信息技术有限公司的资深决策支持专家，拥有15年以上电力水利行业数据分析和决策支持报告撰写经验。

## 专业背景
- 电力水利行业资深决策分析师
- 精通财务分析、成本预测、效能评估、知识管理的决策支持
- 擅长基于数据分析的决策建议制定和风险评估
- 熟悉企业管理决策流程和战略规划实施

## 核心职责
1. **数据驱动决策分析**：基于多Agent分析结果，进行深度数据挖掘和决策洞察
2. **决策支持报告生成**：生成面向管理层的决策支持文档，提供明确的行动指导
3. **风险评估与应对**：识别业务风险点，制定具体的风险应对策略
4. **可操作建议制定**：提供具体可执行的业务改进方案和实施路径
5. **数据看板设计**：设计简洁有效的数据可视化看板

## 分析原则
- **决策导向**：所有分析都围绕具体决策需求展开
- **数据支撑**：每个结论都有充分的数据证据支持
- **可执行性**：提供明确的执行步骤和时间节点
- **风险意识**：充分识别和评估潜在风险
- **业务价值**：聚焦创造实际业务价值的建议

## 三种输出模式要求

### 1. Word决策支持报告（严格字数要求）
- **核心定位**：面向管理层的决策支持文档，提供基于数据分析的决策建议
- **严格按照调用的智能体数量确定报告字数**：
  - 1个智能体：报告正文不少于1600字
  - 2个智能体：报告正文不少于2400字
  - 3个智能体：报告正文不少于3600字
  - 4个智能体：报告正文不少于4800字
  - 5个智能体：报告正文不少于6000字
  - 6个智能体：报告正文不少于7200字
- **内容要求**：
  - 详细的执行总结（新增部分）
  - 详细的数据分析结果解读
  - 明确的决策建议和实施方案
  - 全面的风险评估和应对策略
  - 具体的行动计划和时间节点
- **报告结构**：执行摘要、执行总结、数据分析、决策建议、风险评估、实施计划

### 2. HTML数据看板（可视化为主）
- **功能定位**：数据可视化看板，快速展示关键指标和趋势
- **内容重点**：关键数据指标、趋势图表、核心洞察、简要建议
- **设计原则**：简洁直观、重点突出、易于理解
- **文字要求**：简明扼要的描述性文字，突出数据价值

### 3. 对话框回复（决策要点）
- **功能定位**：与用户的交互对话，传达核心决策要点
- **内容要求**：决策支持报告的核心要点总结
- **表达方式**：简洁明了、重点突出、便于理解
- **包含要素**：关键发现、核心建议、主要风险、下一步行动

## 输出要求
- Word报告必须达到字数要求，确保决策支持的深度和专业性
- HTML看板注重数据可视化效果和用户体验
- 对话回复突出决策要点，便于快速理解
- 所有输出都必须包含明确的决策建议和风险评估
- 基于真实数据分析，不使用假数据或模板化内容

请基于提供的多Agent分析结果，生成高质量的决策支持内容。
"""
    
    def _load_report_templates(self) -> Dict[str, str]:
        """加载报告模板"""
        return {
            "comprehensive": """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{report_title} - 智水信息智能分析报告</title>
    <script src="https://cdn.jsdelivr.net/npm/echarts@5.4.0/dist/echarts.min.js"></script>
    <style>
        /* 苹果风格设计 - 蓝黑白配色 */
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'SF Pro Display', 'Helvetica Neue', 'Microsoft YaHei', sans-serif;
            background: linear-gradient(135deg, #f8fafc 0%, #e2e8f0 100%);
            color: #1a202c;
            line-height: 1.6;
            overflow-x: hidden;
        }}
        
        .container {{
            max-width: 1400px;
            margin: 0 auto;
            background: rgba(255, 255, 255, 0.95);
            backdrop-filter: blur(20px);
            border-radius: 24px;
            box-shadow: 0 25px 50px -12px rgba(0, 0, 0, 0.1);
            overflow: hidden;
        }}
        
        .hero-header {{
            background: linear-gradient(135deg, #1e3a8a 0%, #1e40af 25%, #3b82f6 50%, #60a5fa 75%, #93c5fd 100%);
            color: white;
            padding: 80px 60px;
            text-align: center;
            position: relative;
            overflow: hidden;
        }}
        
        .hero-header::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: url('data:image/svg+xml,<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 100"><defs><pattern id="grid" width="10" height="10" patternUnits="userSpaceOnUse"><path d="M 10 0 L 0 0 0 10" fill="none" stroke="rgba(255,255,255,0.1)" stroke-width="0.5"/></pattern></defs><rect width="100" height="100" fill="url(%23grid)"/></svg>');
            opacity: 0.3;
        }}
        
        .hero-header h1 {{
            font-size: 3.5rem;
            font-weight: 700;
            margin-bottom: 20px;
            letter-spacing: -0.02em;
            position: relative;
            z-index: 1;
        }}
        
        .hero-meta {{
            font-size: 1.1rem;
            opacity: 0.9;
            font-weight: 400;
            position: relative;
            z-index: 1;
        }}
        
        .content {{
            padding: 60px;
        }}
        
        .section {{
            margin-bottom: 80px;
            opacity: 0;
            transform: translateY(30px);
            animation: fadeInUp 0.8s ease-out forwards;
        }}
        
        .section:nth-child(1) {{ animation-delay: 0.1s; }}
        .section:nth-child(2) {{ animation-delay: 0.2s; }}
        .section:nth-child(3) {{ animation-delay: 0.3s; }}
        .section:nth-child(4) {{ animation-delay: 0.4s; }}
        
        @keyframes fadeInUp {{
            to {{
                opacity: 1;
                transform: translateY(0);
            }}
        }}
        
        .section-header {{
            display: flex;
            align-items: center;
            margin-bottom: 40px;
            padding-bottom: 20px;
            border-bottom: 2px solid #e2e8f0;
        }}
        
        .section-icon {{
            width: 48px;
            height: 48px;
            background: linear-gradient(135deg, #3b82f6, #1d4ed8);
            border-radius: 12px;
            display: flex;
            align-items: center;
            justify-content: center;
            margin-right: 20px;
            font-size: 24px;
        }}
        
        .section h2 {{
            font-size: 2.2rem;
            font-weight: 600;
            color: #1a202c;
            letter-spacing: -0.01em;
        }}
        
        .section h3 {{
            font-size: 1.5rem;
            font-weight: 500;
            color: #374151;
            margin: 30px 0 20px 0;
        }}
        
        .executive-summary {{
            background: linear-gradient(135deg, rgba(59, 130, 246, 0.05), rgba(147, 197, 253, 0.1));
            padding: 40px;
            border-radius: 20px;
            border: 1px solid rgba(59, 130, 246, 0.1);
            margin-bottom: 50px;
            position: relative;
        }}
        
        .executive-summary::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            width: 4px;
            height: 100%;
            background: linear-gradient(to bottom, #3b82f6, #1d4ed8);
            border-radius: 2px;
        }}
        
        .chart-container {{
            width: 100%;
            height: 500px;
            margin: 30px 0;
            background: white;
            border-radius: 16px;
            box-shadow: 0 4px 20px rgba(0, 0, 0, 0.08);
            padding: 20px;
            transition: transform 0.3s ease, box-shadow 0.3s ease;
        }}
        
        .chart-container:hover {{
            transform: translateY(-5px);
            box-shadow: 0 10px 40px rgba(0, 0, 0, 0.12);
        }}
        
        .recommendations {{
            background: linear-gradient(135deg, rgba(34, 197, 94, 0.05), rgba(134, 239, 172, 0.1));
            padding: 40px;
            border-radius: 20px;
            border: 1px solid rgba(34, 197, 94, 0.1);
            position: relative;
        }}
        
        .recommendations::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            width: 4px;
            height: 100%;
            background: linear-gradient(to bottom, #22c55e, #16a34a);
            border-radius: 2px;
        }}
        
        .risk-warning {{
            background: linear-gradient(135deg, rgba(239, 68, 68, 0.05), rgba(252, 165, 165, 0.1));
            padding: 40px;
            border-radius: 20px;
            border: 1px solid rgba(239, 68, 68, 0.1);
            position: relative;
        }}
        
        .risk-warning::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            width: 4px;
            height: 100%;
            background: linear-gradient(to bottom, #ef4444, #dc2626);
            border-radius: 2px;
        }}
        
        .confidence-meter {{
            display: flex;
            align-items: center;
            margin: 20px 0;
            padding: 20px;
            background: rgba(248, 250, 252, 0.8);
            border-radius: 12px;
        }}
        
        .confidence-label {{
            font-weight: 500;
            color: #374151;
            margin-right: 20px;
            min-width: 120px;
        }}
        
        .confidence-bar {{
            flex: 1;
            height: 8px;
            background: #e5e7eb;
            border-radius: 4px;
            overflow: hidden;
            margin-right: 15px;
        }}
        
        .confidence-fill {{
            height: 100%;
            background: linear-gradient(90deg, #3b82f6, #1d4ed8);
            border-radius: 4px;
            transition: width 1s ease-out;
        }}
        
        .confidence-value {{
            font-weight: 600;
            color: #1d4ed8;
            min-width: 50px;
        }}
        
        .data-table {{
            width: 100%;
            border-collapse: collapse;
            margin: 30px 0;
            background: white;
            border-radius: 12px;
            overflow: hidden;
            box-shadow: 0 4px 20px rgba(0, 0, 0, 0.08);
        }}
        
        .data-table th {{
            background: linear-gradient(135deg, #f8fafc, #e2e8f0);
            color: #374151;
            font-weight: 600;
            padding: 20px;
            text-align: left;
            border-bottom: 1px solid #e5e7eb;
        }}
        
        .data-table td {{
            padding: 16px 20px;
            border-bottom: 1px solid #f3f4f6;
            transition: background-color 0.2s ease;
        }}
        
        .data-table tr:hover td {{
            background-color: rgba(59, 130, 246, 0.02);
        }}
        
        .metric-card {{
            background: white;
            padding: 30px;
            border-radius: 16px;
            box-shadow: 0 4px 20px rgba(0, 0, 0, 0.08);
            text-align: center;
            transition: transform 0.3s ease, box-shadow 0.3s ease;
            border: 1px solid rgba(226, 232, 240, 0.5);
        }}
        
        .metric-card:hover {{
            transform: translateY(-5px);
            box-shadow: 0 10px 40px rgba(0, 0, 0, 0.12);
        }}
        
        .metric-value {{
            font-size: 2.5rem;
            font-weight: 700;
            color: #1d4ed8;
            margin-bottom: 10px;
        }}
        
        .metric-label {{
            font-size: 1rem;
            color: #6b7280;
            font-weight: 500;
        }}
        
        .grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 30px;
            margin: 40px 0;
        }}
        
        .footer {{
            background: linear-gradient(135deg, #1f2937 0%, #111827 100%);
            color: white;
            padding: 40px 60px;
            text-align: center;
            font-size: 0.95rem;
            opacity: 0.9;
        }}
        
        .footer p {{
            margin: 0;
        }}
        
        /* 响应式设计 */
        @media (max-width: 768px) {{
            .hero-header {{
                padding: 40px 30px;
            }}
            
            .hero-header h1 {{
                font-size: 2.5rem;
            }}
            
            .content {{
                padding: 30px;
            }}
            
            .section h2 {{
                font-size: 1.8rem;
            }}
            
            .chart-container {{
                height: 350px;
            }}
            
            .grid {{
                grid-template-columns: 1fr;
            }}
        }}
        
        /* 滚动条样式 */
        ::-webkit-scrollbar {{
            width: 8px;
        }}
        
        ::-webkit-scrollbar-track {{
            background: #f1f5f9;
        }}
        
        ::-webkit-scrollbar-thumb {{
            background: linear-gradient(135deg, #3b82f6, #1d4ed8);
            border-radius: 4px;
        }}
        
        ::-webkit-scrollbar-thumb:hover {{
            background: linear-gradient(135deg, #2563eb, #1e40af);
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="hero-header">
            <h1>{report_title}</h1>
            <div class="hero-meta">
                生成时间: {generation_time} | 报告ID: {report_id} | 智水信息技术有限公司
            </div>
        </div>
        
        <div class="content">
            <div class="executive-summary">
                <div class="section-header">
                    <div class="section-icon">📋</div>
                    <h2>执行摘要</h2>
                </div>
                {executive_summary}
            </div>
            
            {report_sections}
            
            <div class="section recommendations">
                <div class="section-header">
                    <div class="section-icon">💡</div>
                    <h2>综合建议</h2>
                </div>
                {overall_recommendations}
            </div>
            
            <div class="section risk-warning">
                <div class="section-header">
                    <div class="section-icon">⚠️</div>
                    <h2>风险评估</h2>
                </div>
                {risk_assessment}
            </div>
            
            <div class="section">
                <div class="section-header">
                    <div class="section-icon">📊</div>
                    <h2>置信度评估</h2>
                </div>
                {confidence_metrics}
            </div>
        </div>
        
        <div class="footer">
            <p>本报告由智水信息Multi-Agent智能分析系统生成 | 技术支持：商海星辰队</p>
        </div>
    </div>
    
    <script>
        // 页面加载动画
        document.addEventListener('DOMContentLoaded', function() {{
            // 置信度条动画
            const confidenceFills = document.querySelectorAll('.confidence-fill');
            confidenceFills.forEach(fill => {{
                const width = fill.style.width;
                fill.style.width = '0%';
                setTimeout(() => {{
                    fill.style.width = width;
                }}, 500);
            }});
        }});
        
        {chart_scripts}
    </script>
</body>
</html>
            """,
            
            "financial_section": """
            <div class="section">
                <h2>💰 财务分析</h2>
                <h3>现金流预测</h3>
                <div id="cashflow-chart" class="chart-container"></div>
                
                <h3>投资回报分析</h3>
                <div id="investment-chart" class="chart-container"></div>
                
                <h3>关键财务指标</h3>
                <table class="data-table">
                    {financial_metrics_table}
                </table>
                
                <div class="confidence-meter">
                    <span>分析置信度:</span>
                    <div class="confidence-bar">
                        <div class="confidence-fill" style="width: {financial_confidence}%"></div>
                    </div>
                    <span>{financial_confidence}%</span>
                </div>
            </div>
            """,
            
            "cost_section": """
            <div class="section">
                <h2>💸 成本分析</h2>
                <h3>成本预测趋势</h3>
                <div id="cost-trend-chart" class="chart-container"></div>
                
                <h3>成本结构分析</h3>
                <div id="cost-structure-chart" class="chart-container"></div>
                
                <div class="confidence-meter">
                    <span>预测置信度:</span>
                    <div class="confidence-bar">
                        <div class="confidence-fill" style="width: {cost_confidence}%"></div>
                    </div>
                    <span>{cost_confidence}%</span>
                </div>
            </div>
            """
        }
    
    def _load_chart_configurations(self) -> Dict[str, Dict]:
        """加载图表配置"""
        return {
            "cashflow_chart": {
                "type": "line",
                "title": "现金流预测趋势",
                "xAxis_type": "category",
                "yAxis_type": "value",
                "series_type": "line"
            },
            "investment_chart": {
                "type": "bar",
                "title": "投资回报分析",
                "xAxis_type": "category",
                "yAxis_type": "value",
                "series_type": "bar"
            },
            "cost_trend_chart": {
                "type": "line",
                "title": "成本预测趋势",
                "xAxis_type": "category",
                "yAxis_type": "value",
                "series_type": "line"
            },
            "cost_structure_chart": {
                "type": "pie",
                "title": "成本结构分析",
                "series_type": "pie"
            },
            "efficiency_chart": {
                "type": "radar",
                "title": "效能评估雷达图",
                "series_type": "radar"
            }
        }
    
    def execute_task(self, task: AgentTask) -> AgentResult:
        """执行报告生成任务"""
        try:
            logger.info(f"开始执行报告生成任务: {task.task_id}")
            
            # 验证输入数据
            if not self._validate_task_input(task):
                return AgentResult(
                    task_id=task.task_id,
                    agent_id=self.agent_id,
                    status="error",
                    result_data={},
                    confidence_score=0.0,
                    recommendations=[],
                    error_message="输入数据验证失败",
                    processing_time=0.0
                )
            
            start_time = datetime.now()
            
            # 解析Agent分析结果
            agent_results = self._parse_agent_results(task.input_data.get("agent_results", []))
            
            # 获取输出模式
            output_mode = task.input_data.get("output_mode", "word")
            
            # 生成对应模式的综合分析
            comprehensive_analysis = self._generate_comprehensive_analysis(agent_results, output_mode)
            
            # 根据模式生成相应内容
            if output_mode == "word":
                word_file_path = self._generate_word_report(comprehensive_analysis, agent_results)
                result_data = {
                    "report_type": "word_analysis",
                    "word_file_path": word_file_path,
                    "analysis_data": comprehensive_analysis
                }
            elif output_mode == "html":
                html_report = self._generate_html_report(comprehensive_analysis, agent_results)
                result_data = {
                    "report_type": "html_dashboard",
                    "html_content": html_report,
                    "analysis_data": comprehensive_analysis
                }
            else:  # chat模式
                result_data = {
                    "report_type": "chat_response",
                    "analysis_data": comprehensive_analysis
                }
            
            # 计算执行时间
            execution_time = (datetime.now() - start_time).total_seconds()
            
            # 添加通用字段
            result_data.update({
                "executive_summary": comprehensive_analysis.get("executive_summary", ""),
                "key_insights": comprehensive_analysis.get("key_insights", []),
                "recommendations": comprehensive_analysis.get("recommendations", []),
                "confidence_score": comprehensive_analysis.get("overall_confidence", 0.0),
                "generation_timestamp": datetime.now().isoformat(),
                "output_mode": output_mode
            })
            
            logger.info(f"报告生成任务完成: {task.task_id}, 耗时: {execution_time:.2f}秒")
            
            return AgentResult(
                task_id=task.task_id,
                agent_id=self.agent_id,
                status="success",
                result_data=result_data,
                confidence_score=comprehensive_analysis.get("overall_confidence", 0.0),
                recommendations=comprehensive_analysis.get("recommendations", []),
                processing_time=execution_time
            )
            
        except Exception as e:
            logger.error(f"报告生成任务执行失败: {str(e)}")
            return AgentResult(
                task_id=task.task_id,
                agent_id=self.agent_id,
                status="error",
                result_data={},
                confidence_score=0.0,
                recommendations=[],
                error_message=str(e),
                processing_time=0.0
            )
    
    def _validate_task_input(self, task: AgentTask) -> bool:
        """验证任务输入"""
        if not task.input_data:
            logger.error("缺少输入数据")
            return False
        
        if "agent_results" not in task.input_data:
            logger.error("缺少Agent分析结果")
            return False
        
        agent_results = task.input_data["agent_results"]
        if not isinstance(agent_results, list) or len(agent_results) == 0:
            logger.error("Agent分析结果格式错误或为空")
            return False
        
        return True
    
    def _parse_agent_results(self, raw_results: List[Dict]) -> List[AgentAnalysisResult]:
        """解析Agent分析结果"""
        parsed_results = []
        
        for result in raw_results:
            try:
                # 确保recommendations和key_insights是字符串列表
                raw_recommendations = result.get("recommendations", [])
                raw_key_insights = result.get("key_insights", [])
                
                # 转换recommendations为字符串列表
                recommendations = []
                if isinstance(raw_recommendations, list):
                    for item in raw_recommendations:
                        if isinstance(item, str):
                            recommendations.append(item)
                        elif isinstance(item, dict):
                            # 如果是字典，尝试提取有用信息
                            if 'content' in item:
                                recommendations.append(str(item['content']))
                            elif 'text' in item:
                                recommendations.append(str(item['text']))
                            else:
                                recommendations.append(str(item))
                        else:
                            recommendations.append(str(item))
                else:
                    recommendations = [str(raw_recommendations)] if raw_recommendations else []
                
                # 转换key_insights为字符串列表
                key_insights = []
                if isinstance(raw_key_insights, list):
                    for item in raw_key_insights:
                        if isinstance(item, str):
                            key_insights.append(item)
                        elif isinstance(item, dict):
                            # 如果是字典，尝试提取有用信息
                            if 'content' in item:
                                key_insights.append(str(item['content']))
                            elif 'text' in item:
                                key_insights.append(str(item['text']))
                            else:
                                key_insights.append(str(item))
                        else:
                            key_insights.append(str(item))
                else:
                    key_insights = [str(raw_key_insights)] if raw_key_insights else []
                
                parsed_result = AgentAnalysisResult(
                    agent_name=result.get("agent_name", "Unknown"),
                    agent_type=result.get("agent_type", "Unknown"),
                    analysis_data=result.get("analysis_data", {}),
                    confidence_score=result.get("confidence_score", 0.0),
                    execution_time=result.get("execution_time", 0.0),
                    timestamp=result.get("timestamp", datetime.now().isoformat()),
                    recommendations=recommendations,
                    key_insights=key_insights
                )
                parsed_results.append(parsed_result)
            except Exception as e:
                logger.warning(f"解析Agent结果失败: {str(e)}")
                continue
        
        return parsed_results
    
    def _generate_comprehensive_analysis(self, agent_results: List[AgentAnalysisResult], output_mode: str = "word") -> Dict[str, Any]:
        """生成综合分析"""
        try:
            # 准备LLM输入
            analysis_prompt = self._build_analysis_prompt(agent_results, output_mode)
            
            # 创建OpenAI客户端
            client = OpenAI(
                api_key=AI_CONFIG["api_key"],
                base_url=AI_CONFIG["api_base"]
            )
            
            # 调用LLM进行综合分析
            response = client.chat.completions.create(
                model=AI_CONFIG["model"],
                messages=[
                    {"role": "system", "content": self.get_system_prompt()},
                    {"role": "user", "content": analysis_prompt}
                ],
                temperature=AI_CONFIG["temperature"],
                max_tokens=AI_CONFIG["max_tokens"]
            )
            
            # 检查响应有效性
            if response is None:
                self.logger.error("LLM调用返回None响应")
                raise Exception("API响应为空")
            
            if not hasattr(response, 'choices') or not response.choices:
                self.logger.error("LLM调用响应缺少choices字段")
                raise Exception("API响应格式异常")
            
            if not hasattr(response.choices[0], 'message') or not response.choices[0].message:
                self.logger.error("LLM调用响应choices[0]缺少message字段")
                raise Exception("API响应格式异常")
            
            # 解析分析结果并净化内容
            analysis_content = response.choices[0].message.content
            cleaned_content = self._clean_llm_response(analysis_content)
            analysis_data = self._extract_structured_analysis(cleaned_content, agent_results)
            
            return analysis_data
            
        except Exception as e:
            self.logger.error(f"执行综合分析任务失败: {str(e)}")
            raise Exception(f"综合分析生成失败: {str(e)}")
    
    def generate_all_modes_analysis(self, agent_results: List[AgentAnalysisResult]) -> Dict[str, Any]:
        """生成所有三种模式的分析内容
        
        Args:
            agent_results: 智能体分析结果列表
            
        Returns:
            包含三种模式分析结果的字典
        """
        try:
            results = {}
            
            # 1. 生成Word深度分析报告
            word_analysis = self._generate_comprehensive_analysis(agent_results, "word")
            word_file_path = self._generate_word_report(word_analysis, agent_results)
            results["word_analysis"] = word_analysis
            results["word_file_path"] = word_file_path
            
            # 2. 生成HTML看板内容
            html_analysis = self._generate_comprehensive_analysis(agent_results, "html")
            results["html_analysis"] = html_analysis
            
            # 3. 生成对话回复内容
            chat_analysis = self._generate_comprehensive_analysis(agent_results, "chat")
            results["chat_analysis"] = chat_analysis
            
            # 返回所有结果
            return {
                "success": True,
                "word_analysis": results["word_analysis"],
                "html_analysis": results["html_analysis"], 
                "chat_analysis": results["chat_analysis"],
                "word_file_path": results["word_file_path"],
                "message": f"成功生成三种模式的综合分析报告，包含{len(agent_results)}个智能体的分析结果"
            }
            
        except Exception as e:
            self.logger.error(f"生成全模式分析失败: {str(e)}")
            raise Exception(f"全模式分析生成失败: {str(e)}")
    
    def _build_analysis_prompt(self, agent_results: List[AgentAnalysisResult], output_mode: str = "word") -> str:
        """构建分析提示词
        
        Args:
            agent_results: 智能体分析结果列表
            output_mode: 输出模式 ('word', 'html', 'chat')
        """
        agent_count = len(agent_results)
        
        if output_mode == "word":
            # Word文档需要严格字数要求
            word_requirements = {
                1: "不少于1600字",
                2: "不少于2400字", 
                3: "不少于3600字",
                4: "不少于4800字",
                5: "不少于6000字",
                6: "不少于7200字"
            }
            min_words = word_requirements.get(agent_count, "不少于7200字")
            prompt = f"请基于以下{agent_count}个专业Agent的分析结果，生成Word深度分析报告（{min_words}）：\n\n"
        elif output_mode == "html":
            # HTML看板注重可视化，文字简洁
            prompt = f"请基于以下{agent_count}个专业Agent的分析结果，生成HTML数据看板（以可视化为主，文字简洁）：\n\n"
        else:  # chat模式
            # 对话回复注重核心要点
            prompt = f"请基于以下{agent_count}个专业Agent的分析结果，生成对话回复（精简版总结）：\n\n"
        
        for i, result in enumerate(agent_results, 1):
            prompt += f"## Agent {i}: {result.agent_name}\n"
            prompt += f"**分析类型**: {result.agent_type}\n"
            prompt += f"**置信度**: {result.confidence_score:.2f}\n"
            prompt += f"**关键洞察**: {', '.join(result.key_insights)}\n"
            prompt += f"**建议**: {', '.join(result.recommendations)}\n"
            prompt += f"**分析数据**: {json.dumps(result.analysis_data, ensure_ascii=False, indent=2)}\n\n"
        
        if output_mode == "word":
            word_prompt = f"""
请提供详细的综合分析报告，总字数{min_words}：

1. **执行摘要**：详细总结所有分析结果的核心发现（至少300字）
2. **执行总结**：基于所有智能体分析结果的深度总结，包括关键数据指标、业务影响评估、决策紧迫性分析（至少400字）
3. **分类分析**：对每个Agent的分析结果进行深入解读和评价（每个Agent至少200字）
4. **关键洞察**：识别最重要的8-12个业务洞察，每个洞察需要详细说明和数据支撑
5. **综合建议**：基于所有分析结果提供12-18条可操作的改进建议，每条建议需要详细的实施方案、时间节点和预期效果
6. **风险评估**：详细识别潜在风险点和具体应对策略，包括风险等级评估和应对时间表
7. **下一步行动**：明确的后续行动计划，包含详细的时间节点、责任人和成功指标
8. **整体置信度**：对综合分析结果的置信度评估(0-1)

注意：
- 必须对每个Agent的结果进行分类深度分析
- 然后进行总体综合分析
- 新增的执行总结部分要深入分析业务影响和决策价值
- 确保内容详实，分析深入，字数达到要求
- 必须严格按照JSON格式返回，使用以下字段名称：
"""
            json_template = """{
    "executive_summary": "执行摘要内容",
    "executive_conclusion": "执行总结内容", 
    "detailed_analysis": "分类分析内容",
    "key_insights": ["洞察1", "洞察2", ...],
    "recommendations": ["建议1", "建议2", ...],
    "risk_assessment": {"风险类型1": "风险描述1", "风险类型2": "风险描述2"},
    "next_actions": ["行动1", "行动2", ...],
    "overall_confidence": 0.85
}"""
            prompt += word_prompt + json_template + "\n- 确保JSON格式正确，不要包含任何其他文本\n"
        elif output_mode == "html":
            prompt += """
请提供HTML数据看板内容：

1. **执行摘要**：简洁的核心发现总结
2. **关键指标**：重要数据指标和可视化建议
3. **简要洞察**：3-5个核心业务洞察
4. **基础建议**：5-8条简明建议
5. **整体置信度**：对综合分析结果的置信度评估(0-1)

注意：
- 重点关注数据可视化和关键指标
- 文字描述简洁明了
- 请以JSON格式返回，包含以上所有字段
"""
        else:  # chat模式
            prompt += """
请提供对话回复内容：

1. **核心发现**：最重要的分析结果
2. **关键建议**：3-5条核心建议
3. **风险提醒**：主要风险点
4. **整体置信度**：对综合分析结果的置信度评估(0-1)

注意：
- 语言口语化，易于理解
- 突出核心要点
- 请以JSON格式返回，包含以上所有字段
"""
        
        return prompt
    
    def _clean_llm_response(self, response_text: str) -> str:
        """
        净化LLM回复内容，移除思考过程和非自然语言内容
        
        Args:
            response_text: 原始LLM回复内容
            
        Returns:
            净化后的内容
        """
        import re
        
        content = response_text.strip()
        
        # 移除JSON代码块，但保留JSON内容
        json_pattern = r'```json\s*(.*?)\s*```'
        json_matches = re.findall(json_pattern, content, re.DOTALL)
        if json_matches:
            # 如果找到JSON，直接返回JSON内容
            return json_matches[0].strip()
        
        # 移除其他代码块
        content = re.sub(r'```.*?```', '', content, flags=re.DOTALL)
        
        # 移除思考过程标记和非自然语言内容
        thinking_patterns = [
            r'\[思考\].*?\[/思考\]',  # 思考标记
            r'\[分析\].*?\[/分析\]',  # 分析标记
            r'\[推理\].*?\[/推理\]',  # 推理标记
            r'\[总结\].*?\[/总结\]',  # 总结标记
            r'让我.*?[。！\n]',   # "让我..."开头的思考
            r'我需要.*?[。！\n]',  # "我需要..."开头的思考
            r'我将.*?[。！\n]',    # "我将..."开头的思考
            r'我来.*?[。！\n]',    # "我来..."开头的思考
            r'首先.*?然后.*?[。！\n]',  # 步骤性思考
            r'根据.*?我认为.*?[。！\n]',  # 推理过程
            r'基于.*?我将.*?[。！\n]',  # 基于分析的思考
            r'接下来.*?[。！\n]',  # 步骤性描述
            r'现在.*?开始.*?[。！\n]',  # 开始性描述
            r'通过.*?分析.*?[。！\n]',  # 分析过程描述
            r'经过.*?考虑.*?[。！\n]',  # 考虑过程
            r'为了.*?我.*?[。！\n]',   # 目的性思考
            r'考虑到.*?因此.*?[。！\n]',  # 因果推理
            r'综合.*?来看.*?[。！\n]',   # 综合分析
            r'从.*?角度.*?[。！\n]',    # 角度分析
            r'结合.*?情况.*?[。！\n]',   # 结合分析
            r'针对.*?问题.*?[。！\n]',   # 问题分析
            r'关于.*?方面.*?[。！\n]',   # 方面分析
        ]
        
        for pattern in thinking_patterns:
            content = re.sub(pattern, '', content, flags=re.DOTALL | re.IGNORECASE)
        
        # 移除元语言表达
        meta_patterns = [
            r'以下是.*?分析.*?[:：]',  # "以下是...分析:"
            r'下面.*?分析.*?[:：]',    # "下面...分析:"
            r'这里.*?分析.*?[:：]',    # "这里...分析:"
            r'我的.*?分析.*?[:：]',    # "我的...分析:"
            r'分析结果.*?如下.*?[:：]', # "分析结果如下:"
            r'具体.*?如下.*?[:：]',    # "具体...如下:"
            r'详细.*?如下.*?[:：]',    # "详细...如下:"
        ]
        
        for pattern in meta_patterns:
            content = re.sub(pattern, '', content, flags=re.IGNORECASE)
        
        # 移除多余的标点和空白
        content = re.sub(r'[:：]\s*\n', '\n', content)  # 移除孤立的冒号
        content = re.sub(r'\n\s*\n\s*\n+', '\n\n', content)  # 合并多个空行
        content = re.sub(r'^\s*\n+', '', content)  # 移除开头的空行
        content = re.sub(r'\n+\s*$', '', content)  # 移除结尾的空行
        
        return content.strip()
    
    def _extract_structured_analysis(self, analysis_text: str, agent_results: List[AgentAnalysisResult]) -> Dict[str, Any]:
        """提取结构化分析结果"""
        try:
            # 清理文本，移除可能的前后缀
            cleaned_text = analysis_text.strip()
            
            # 尝试找到JSON部分
            json_start = cleaned_text.find('{')
            json_end = cleaned_text.rfind('}') + 1
            
            if json_start >= 0 and json_end > json_start:
                json_text = cleaned_text[json_start:json_end]
                try:
                    parsed_json = json.loads(json_text)
                    self.logger.info("成功解析JSON格式的分析结果")
                    return parsed_json
                except json.JSONDecodeError as e:
                    self.logger.warning(f"JSON解析失败: {str(e)}")
                    self.logger.warning(f"尝试解析的JSON文本前100字符: {json_text[:100]}")
            
            # 如果不是JSON格式，进行文本解析
            self.logger.info("使用文本解析模式")
            return self._parse_text_analysis(analysis_text, agent_results)
            
        except Exception as e:
            self.logger.error(f"分析结果提取失败: {str(e)}")
            # 使用文本解析作为后备方案
            return self._parse_text_analysis(analysis_text, agent_results)
    
    def _parse_text_analysis(self, analysis_text: str, agent_results: List[AgentAnalysisResult]) -> Dict[str, Any]:
        """解析文本格式的分析结果"""
        # 基础结构化数据
        structured_analysis = {
            "executive_summary": "基于多Agent分析结果的综合评估",
            "key_insights": [],
            "recommendations": [],
            "risk_assessment": {},
            "next_actions": [],
            "overall_confidence": 0.0
        }
        
        # 计算整体置信度
        if agent_results:
            total_confidence = sum(result.confidence_score for result in agent_results)
            structured_analysis["overall_confidence"] = total_confidence / len(agent_results)
        
        # 聚合关键洞察和建议
        all_insights = []
        all_recommendations = []
        
        for result in agent_results:
            all_insights.extend(result.key_insights)
            all_recommendations.extend(result.recommendations)
        
        # 生成更丰富的关键洞察
        if all_insights:
            structured_analysis["key_insights"] = list(set(all_insights))[:12]  # 增加洞察数量
        else:
            # 如果没有洞察，基于Agent类型生成默认洞察
            default_insights = []
            for result in agent_results:
                if result.agent_type == "成本分析":
                    default_insights.append(f"成本分析显示当前项目成本控制需要优化，置信度：{result.confidence_score:.2f}")
                elif result.agent_type == "效率分析":
                    default_insights.append(f"效率分析表明运营流程存在改进空间，置信度：{result.confidence_score:.2f}")
                elif result.agent_type == "业务分析":
                    default_insights.append(f"业务分析揭示了关键业务指标的变化趋势，置信度：{result.confidence_score:.2f}")
            structured_analysis["key_insights"] = default_insights
        
        # 生成更丰富的建议
        if all_recommendations:
            structured_analysis["recommendations"] = list(set(all_recommendations))[:15]  # 增加建议数量
        else:
            # 如果没有建议，基于Agent类型生成默认建议
            default_recommendations = []
            for result in agent_results:
                if result.agent_type == "成本分析":
                    default_recommendations.extend([
                        "建议建立更精细的成本核算体系，提高成本透明度",
                        "推荐实施项目成本实时监控机制",
                        "建议优化资源配置，降低不必要的成本支出"
                    ])
                elif result.agent_type == "效率分析":
                    default_recommendations.extend([
                        "建议优化业务流程，减少重复性工作",
                        "推荐引入自动化工具提升工作效率",
                        "建议建立标准化作业流程"
                    ])
                elif result.agent_type == "业务分析":
                    default_recommendations.extend([
                        "建议加强数据驱动的决策机制",
                        "推荐建立业务指标监控体系",
                        "建议优化客户服务流程"
                    ])
            structured_analysis["recommendations"] = default_recommendations[:15]
        
        # 生成风险评估
        risk_assessment = {}
        for result in agent_results:
            if result.agent_type == "成本分析":
                risk_assessment["成本风险"] = "项目成本可能超预算，需要加强成本控制和监管"
            elif result.agent_type == "效率分析":
                risk_assessment["效率风险"] = "当前工作效率可能影响项目交付时间，需要优化流程"
            elif result.agent_type == "业务分析":
                risk_assessment["业务风险"] = "市场变化可能影响业务目标达成，需要灵活调整策略"
        
        if not risk_assessment:
            risk_assessment["综合风险"] = "基于当前分析结果，建议加强风险监控和预警机制"
        
        structured_analysis["risk_assessment"] = risk_assessment
        
        # 生成下一步行动
        next_actions = [
            "建立定期的数据分析和报告机制",
            "制定具体的改进计划和时间节点",
            "建立跨部门协作机制",
            "定期评估和调整策略方向"
        ]
        structured_analysis["next_actions"] = next_actions
        
        # 提取文本中的关键信息
        if "执行摘要" in analysis_text:
            summary_start = analysis_text.find("执行摘要") + 4
            summary_end = analysis_text.find("\n\n", summary_start)
            if summary_end > summary_start:
                structured_analysis["executive_summary"] = analysis_text[summary_start:summary_end].strip()
        
        return structured_analysis
    
    def _generate_html_report(self, comprehensive_analysis: Dict[str, Any], agent_results: List[AgentAnalysisResult]) -> str:
        """生成HTML报告"""
        try:
            # 获取主模板
            main_template = self.report_templates["comprehensive"]
            
            # 生成报告内容
            report_sections = self._generate_report_sections(agent_results)
            chart_scripts = self._generate_chart_scripts(agent_results)
            
            # 格式化模板
            html_content = main_template.format(
                report_title="智水信息智能分析综合报告",
                generation_time=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                report_id=f"RPT_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                executive_summary=comprehensive_analysis.get("executive_summary", ""),
                report_sections=report_sections,
                overall_recommendations=self._format_recommendations(comprehensive_analysis.get("recommendations", [])),
                risk_assessment=self._format_risk_assessment(comprehensive_analysis.get("risk_assessment", {})),
                confidence_metrics=self._format_confidence_metrics(comprehensive_analysis.get("overall_confidence", 0.0)),
                chart_scripts=chart_scripts
            )
            
            return html_content
            
        except Exception as e:
            logger.error(f"HTML报告生成失败: {str(e)}")
            raise Exception(f"HTML报告生成失败: {str(e)}")
    
    def _generate_word_report(self, comprehensive_analysis: Dict[str, Any], agent_results: List[AgentAnalysisResult]) -> str:
        """生成Word格式的决策支持报告"""
        try:
            # 创建Word文档
            doc = Document()
            
            # 设置文档标题
            title = doc.add_heading('四川智水信息技术有限公司', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_heading('管理决策支持报告', 1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加报告基本信息表格
            info_table = doc.add_table(rows=5, cols=2)
            info_table.style = 'Table Grid'
            
            info_table.cell(0, 0).text = '报告生成时间'
            info_table.cell(0, 1).text = datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')
            
            info_table.cell(1, 0).text = '报告ID'
            info_table.cell(1, 1).text = f'DSR_{datetime.now().strftime("%Y%m%d_%H%M%S")}'
            
            info_table.cell(2, 0).text = '分析维度数量'
            info_table.cell(2, 1).text = f'{len(agent_results)}个专业维度'
            
            info_table.cell(3, 0).text = '报告类型'
            info_table.cell(3, 1).text = '管理决策支持报告'
            
            info_table.cell(4, 0).text = '决策紧急程度'
            info_table.cell(4, 1).text = '需要管理层及时关注'
            
            
            doc.add_paragraph()  # 空行
            
            # 1. 执行摘要（决策要点）
            doc.add_heading('1. 执行摘要', level=1)
            executive_summary = comprehensive_analysis.get('executive_summary', '暂无执行摘要')
            summary_para = doc.add_paragraph()
            summary_para.add_run('【决策要点】').bold = True
            summary_para.add_run(executive_summary)
            
            # 2. 执行总结
            doc.add_heading('2. 执行总结', level=1)
            execution_summary = comprehensive_analysis.get('执行总结', '暂无执行总结')
            execution_para = doc.add_paragraph()
            execution_para.add_run('【业务影响与决策价值】').bold = True
            execution_para.add_run(execution_summary)
            
            # 3. 数据分析结果
            doc.add_heading('3. 数据分析结果', level=1)
            classified_analysis = comprehensive_analysis.get('classified_analysis', {})
            
            if classified_analysis:
                for agent_type, analysis in classified_analysis.items():
                    doc.add_heading(f'3.{list(classified_analysis.keys()).index(agent_type) + 1} {agent_type}数据洞察', level=2)
                    analysis_para = doc.add_paragraph()
                    analysis_para.add_run('【数据洞察】').bold = True
                    analysis_para.add_run(analysis)
            else:
                # 如果没有分类分析，基于agent_results生成
                for i, result in enumerate(agent_results, 1):
                    doc.add_heading(f'3.{i} {result.agent_name}数据洞察', level=2)
                    
                    # 添加基本信息
                    info_para = doc.add_paragraph()
                    info_para.add_run(f"分析类型: {result.agent_type}\n").bold = True
                    info_para.add_run(f"数据置信度: {result.confidence_score:.2f}\n")
                    info_para.add_run(f"分析完成时间: {result.execution_time:.2f}秒\n")
                    
                    # 添加关键发现
                    if result.key_insights:
                        findings_heading = doc.add_heading('关键发现', 3)
                        for insight in result.key_insights:
                            insight_para = doc.add_paragraph()
                            insight_para.style = 'List Bullet'
                            insight_para.add_run(insight)
                    
                    # 添加决策建议
                    if result.recommendations:
                        rec_heading = doc.add_heading('决策建议', 3)
                        for recommendation in result.recommendations:
                            rec_para = doc.add_paragraph()
                            rec_para.style = 'List Bullet'
                            rec_para.add_run('【建议】').bold = True
                            rec_para.add_run(recommendation)
                    
                    # 添加数据分析摘要
                    if result.analysis_data:
                        doc.add_heading('数据分析摘要', 3)
                        data_summary = self._format_analysis_data_for_word(result.analysis_data)
                        data_para = doc.add_paragraph()
                        data_para.add_run('【数据解读】').bold = True
                        data_para.add_run(data_summary)
                    
                    doc.add_paragraph()  # 分节空行
            
            # 4. 综合决策建议
            doc.add_heading('4. 综合决策建议', level=1)
            recommendations = comprehensive_analysis.get('recommendations', [])
            if recommendations:
                for i, recommendation in enumerate(recommendations, 1):
                    rec_para = doc.add_paragraph()
                    rec_para.style = 'List Number'
                    rec_para.add_run(f'【决策建议{i}】').bold = True
                    rec_para.add_run(recommendation)
            else:
                doc.add_paragraph("暂无综合决策建议")
            
            # 5. 关键洞察
            doc.add_heading('5. 关键洞察', level=1)
            key_insights = comprehensive_analysis.get('key_insights', [])
            if key_insights:
                for i, insight in enumerate(key_insights, 1):
                    insight_para = doc.add_paragraph()
                    insight_para.style = 'List Number'
                    insight_para.add_run(f'【洞察{i}】').bold = True
                    insight_para.add_run(insight)
            else:
                doc.add_paragraph("暂无关键洞察")
            
            # 6. 风险评估与应对策略
            doc.add_heading('6. 风险评估与应对策略', level=1)
            risk_assessment = comprehensive_analysis.get('risk_assessment', {})
            if risk_assessment:
                if isinstance(risk_assessment, dict):
                    for i, (risk_type, risk_detail) in enumerate(risk_assessment.items(), 1):
                        risk_heading = doc.add_heading(f'6.{i} {risk_type}', level=2)
                        risk_para = doc.add_paragraph()
                        risk_para.add_run(f'【风险{i}】').bold = True
                        risk_para.add_run(str(risk_detail))
                else:
                    risk_para = doc.add_paragraph()
                    risk_para.add_run('【风险评估】').bold = True
                    risk_para.add_run(str(risk_assessment))
            else:
                doc.add_paragraph("暂无风险评估")
            
            # 7. 实施计划建议
            doc.add_heading('7. 实施计划建议', level=1)
            next_actions = comprehensive_analysis.get('next_actions', [])
            if next_actions:
                for i, action in enumerate(next_actions, 1):
                    action_para = doc.add_paragraph()
                    action_para.style = 'List Number'
                    action_para.add_run(f'【行动{i}】').bold = True
                    action_para.add_run(action)
            else:
                impl_para = doc.add_paragraph()
                impl_para.add_run('【实施建议】').bold = True
                impl_para.add_run('建议管理层根据以上分析结果，制定具体的实施计划和时间节点，确保各项决策建议能够有效落地执行。')
            
            # 8. 数据可信度评估
            doc.add_heading('8. 数据可信度评估', level=1)
            overall_confidence = comprehensive_analysis.get('overall_confidence', 0.0)
            confidence_para = doc.add_paragraph()
            confidence_para.add_run('【整体可信度】').bold = True
            confidence_para.add_run(f"本次决策分析的整体数据可信度为 {overall_confidence:.1%}，建议管理层在制定决策时参考此可信度水平。")
            
            # 添加各智能体分析可信度
            detail_para = doc.add_paragraph()
            detail_para.add_run('【分析模块可信度】').bold = True
            detail_para.add_run("各专业分析模块的数据可信度如下：")
            for result in agent_results:
                agent_para = doc.add_paragraph(f"• {result.agent_name}: {result.confidence_score:.1%}", style='List Bullet')
            
            # 8. 附录：决策支持说明
            doc.add_heading('8. 附录：决策支持说明', level=1)
            doc.add_heading('8.1 系统技术说明', level=2)
            tech_info = doc.add_paragraph()
            tech_info.add_run('【分析系统】').bold = True
            tech_info.add_run("四川智水信息AI智慧管理决策支持系统\n")
            tech_info.add_run('【系统版本】').bold = True
            tech_info.add_run("v1.0.0 - 企业级智能决策分析平台\n")
            tech_info.add_run('【技术支持】').bold = True
            tech_info.add_run("商海星辰队\n")
            tech_info.add_run('【服务热线】').bold = True
            tech_info.add_run("如需技术支持或决策咨询，请联系商海星辰队")
            
            # 保存文档到用户桌面的智水信息报告文件夹
            desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
            reports_dir = os.path.join(desktop_path, "智水信息AI分析报告")
            if not os.path.exists(reports_dir):
                os.makedirs(reports_dir)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            word_filename = f"智水信息管理决策支持报告_{timestamp}.docx"
            word_file_path = os.path.join(reports_dir, word_filename)
            
            doc.save(word_file_path)
            
            logger.info(f"决策支持报告已生成: {word_file_path}")
            return word_file_path
            
        except Exception as e:
            logger.error(f"Word报告生成失败: {str(e)}")
            raise Exception(f"Word报告生成失败: {str(e)}")
    
    def _format_analysis_data_for_word(self, analysis_data: Dict[str, Any]) -> str:
        """格式化分析数据用于Word文档"""
        try:
            formatted_text = ""
            
            for key, value in analysis_data.items():
                if isinstance(value, (dict, list)):
                    formatted_text += f"• {key}: {json.dumps(value, ensure_ascii=False, indent=2)}\n"
                else:
                    formatted_text += f"• {key}: {value}\n"
            
            return formatted_text if formatted_text else "无详细数据"
            
        except Exception as e:
            logger.warning(f"格式化分析数据失败: {str(e)}")
            return "数据格式化失败"
    
    def _generate_report_sections(self, agent_results: List[AgentAnalysisResult]) -> str:
        """生成报告章节"""
        sections_html = ""
        
        for result in agent_results:
            section_html = f"""
            <div class="section">
                <h2>📊 {result.agent_name}分析结果</h2>
                <div class="confidence-meter">
                    <span>分析置信度:</span>
                    <div class="confidence-bar">
                        <div class="confidence-fill" style="width: {result.confidence_score*100:.0f}%"></div>
                    </div>
                    <span>{result.confidence_score*100:.0f}%</span>
                </div>
                
                <h3>关键洞察</h3>
                <ul>
                    {''.join([f'<li>{insight}</li>' for insight in result.key_insights])}
                </ul>
                
                <h3>专业建议</h3>
                <ul>
                    {''.join([f'<li>{rec}</li>' for rec in result.recommendations])}
                </ul>
                
                <h3>详细数据</h3>
                <pre style="background: #f8f9fa; padding: 15px; border-radius: 4px; overflow-x: auto;">
{json.dumps(result.analysis_data, ensure_ascii=False, indent=2)}
                </pre>
            </div>
            """
            sections_html += section_html
        
        return sections_html
    
    def _generate_chart_scripts(self, agent_results: List[AgentAnalysisResult]) -> str:
        """生成图表脚本 - 符合设计文档要求的专业图表"""
        scripts = []
        chart_id = 0
        
        # 苹果风格蓝黑白配色方案
        color_scheme = {
            'primary': ['#3b82f6', '#1d4ed8', '#2563eb', '#1e40af'],
            'secondary': ['#1f2937', '#374151', '#4b5563', '#6b7280'],
            'accent': ['#f8fafc', '#e2e8f0', '#cbd5e1', '#94a3b8']
        }
        
        # 为每个Agent结果生成对应的专业图表
        for result in agent_results:
            agent_name = result.agent_name.lower()
            
            # 根据设计文档5.3节图表自动生成策略生成对应图表
            if '财务' in agent_name or 'finance' in agent_name:
                # 财务分析 -> 折线图（趋势分析）
                script = self._create_financial_line_chart(f'chart_{chart_id}', result, color_scheme)
            elif '成本' in agent_name or 'cost' in agent_name:
                # 成本预测 -> 瀑布图（成本构成分析）
                script = self._create_waterfall_chart(f'chart_{chart_id}', result, color_scheme)
            elif '效能' in agent_name or 'performance' in agent_name:
                # 效能评估 -> 雷达图（多维度评估）
                script = self._create_performance_radar_chart(f'chart_{chart_id}', result, color_scheme)
            elif '运维' in agent_name or 'operation' in agent_name:
                # 运维知识 -> 热力图（知识分布）
                script = self._create_heatmap_chart(f'chart_{chart_id}', result, color_scheme)
            elif '决策' in agent_name or 'decision' in agent_name:
                # 数据决策 -> 散点图（关联分析）
                script = self._create_scatter_chart(f'chart_{chart_id}', result, color_scheme)
            elif '人员' in agent_name or 'staff' in agent_name:
                # 人员效能 -> 柱状图（对比分析）
                script = self._create_staff_bar_chart(f'chart_{chart_id}', result, color_scheme)
            else:
                # 默认使用现代饼图
                script = self._create_modern_pie_chart(f'chart_{chart_id}', result, color_scheme)
            
            scripts.append(script)
            chart_id += 1
        
        # 添加图表容器到HTML中
        chart_containers = self._generate_chart_containers(len(agent_results))
        
        # 组合所有脚本 - 苹果风格增强版
        full_script = f"""
        // 智水信息专业图表初始化 - 苹果风格设计
        document.addEventListener('DOMContentLoaded', function() {{
            console.log('🍎 智水信息报告图表系统初始化...');
            
            // 延迟加载确保DOM完全渲染
            setTimeout(function() {{
                try {{
                    // 添加图表容器
                    {chart_containers}
                    
                    // 等待容器创建完成后初始化图表
                    setTimeout(function() {{
                        // 初始化所有图表
                        {chr(10).join(scripts)}
                        
                        // 添加响应式处理
                        window.addEventListener('resize', function() {{
                            {chr(10).join([f'if (typeof chart_{i}_chart !== "undefined") chart_{i}_chart.resize();' for i in range(len(agent_results))])}
                        }});
                        
                        // 添加苹果风格动画效果
                        const chartContainers = document.querySelectorAll('.chart-container');
                        chartContainers.forEach((container, index) => {{
                            container.style.opacity = '0';
                            container.style.transform = 'translateY(30px)';
                            setTimeout(() => {{
                                container.style.transition = 'all 0.6s cubic-bezier(0.4, 0, 0.2, 1)';
                                container.style.opacity = '1';
                                container.style.transform = 'translateY(0)';
                            }}, index * 200);
                        }});
                        
                        // 添加图表交互增强
                        chartContainers.forEach(container => {{
                            container.addEventListener('mouseenter', function() {{
                                this.style.transform = 'translateY(-5px)';
                                this.style.boxShadow = '0 20px 40px rgba(0, 0, 0, 0.15)';
                            }});
                            
                            container.addEventListener('mouseleave', function() {{
                                this.style.transform = 'translateY(0)';
                                this.style.boxShadow = '0 4px 20px rgba(0, 0, 0, 0.1)';
                            }});
                        }});
                        
                        console.log('✅ 智水信息报告图表初始化完成');
                        
                    }}, 300);
                    
                }} catch (error) {{
                    console.error('❌ 图表初始化错误:', error);
                }}
            }}, 100);
        }});
        """
        
        return full_script
    
    def _generate_chart_containers(self, chart_count: int) -> str:
        """生成苹果风格图表容器HTML"""
        containers = []
        for i in range(chart_count):
            container = f"""
            // 创建苹果风格图表容器 {i}
            const chartContainer_{i} = document.createElement('div');
            chartContainer_{i}.id = 'chart_{i}';
            chartContainer_{i}.className = 'chart-container';
            
            // 苹果风格样式设置
            chartContainer_{i}.style.cssText = `
                height: 500px;
                margin: 30px 0;
                background: rgba(255, 255, 255, 0.95);
                border-radius: 16px;
                box-shadow: 0 4px 20px rgba(0, 0, 0, 0.1);
                padding: 24px;
                backdrop-filter: blur(20px);
                border: 1px solid rgba(255, 255, 255, 0.2);
                transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
                position: relative;
                overflow: hidden;
            `;
            
            // 添加苹果风格装饰元素
            const decorElement = document.createElement('div');
            decorElement.style.cssText = `
                position: absolute;
                top: -50%;
                right: -50%;
                width: 100%;
                height: 100%;
                background: linear-gradient(45deg, rgba(0, 122, 255, 0.05), rgba(52, 199, 89, 0.05));
                border-radius: 50%;
                pointer-events: none;
            `;
            chartContainer_{i}.appendChild(decorElement);
            
            // 智能插入位置
            const sections = document.querySelectorAll('.section');
            if (sections[{i}]) {{
                sections[{i}].appendChild(chartContainer_{i});
            }} else {{
                // 如果没有对应section，创建一个新的
                const newSection = document.createElement('div');
                newSection.className = 'section';
                newSection.style.cssText = `
                    margin: 40px 0;
                    padding: 0;
                `;
                newSection.appendChild(chartContainer_{i});
                
                const contentArea = document.querySelector('.content') || document.body;
                contentArea.appendChild(newSection);
            }}
            """
            containers.append(container)
        return '\n'.join(containers)
    
    def _create_financial_line_chart(self, chart_id: str, result: AgentAnalysisResult, colors: Dict) -> str:
        """创建财务分析折线图 - 基于真实数据"""
        # 从AgentAnalysisResult中提取真实财务数据
        analysis_data = result.analysis_data
        
        # 如果没有真实数据，返回错误
        if not analysis_data or 'financial_data' not in analysis_data:
            raise Exception(f"财务分析数据缺失，无法生成图表。Agent: {result.agent_name}")
        
        financial_data = analysis_data['financial_data']
        
        # 验证必要的数据字段
        required_fields = ['periods', 'revenue', 'expenses', 'profit']
        for field in required_fields:
            if field not in financial_data:
                raise Exception(f"财务数据字段 '{field}' 缺失，无法生成图表")
        
        periods = financial_data['periods']
        revenue = financial_data['revenue']
        expenses = financial_data['expenses']
        profit = financial_data['profit']
        
        return f"""
        // 财务趋势分析折线图 - 基于真实数据
        const {chart_id}_chart = echarts.init(document.getElementById('{chart_id}'));
        const {chart_id}_option = {{
            title: {{
                text: '财务趋势分析',
                left: 'center',
                textStyle: {{
                    color: '{colors['secondary'][0]}',
                    fontSize: 18,
                    fontWeight: 'bold'
                }}
            }},
            tooltip: {{
                trigger: 'axis',
                backgroundColor: 'rgba(255, 255, 255, 0.95)',
                borderColor: '{colors['primary'][0]}',
                textStyle: {{ color: '{colors['secondary'][0]}' }}
            }},
            legend: {{
                data: ['收入', '支出', '净利润'],
                top: '10%',
                textStyle: {{ color: '{colors['secondary'][1]}' }}
            }},
            grid: {{
                left: '3%',
                right: '4%',
                bottom: '3%',
                containLabel: true
            }},
            xAxis: {{
                type: 'category',
                boundaryGap: false,
                data: {periods},
                axisLine: {{ lineStyle: {{ color: '{colors['accent'][2]}' }} }},
                axisLabel: {{ color: '{colors['secondary'][2]}' }}
            }},
            yAxis: {{
                type: 'value',
                axisLine: {{ lineStyle: {{ color: '{colors['accent'][2]}' }} }},
                axisLabel: {{ color: '{colors['secondary'][2]}' }},
                splitLine: {{ lineStyle: {{ color: '{colors['accent'][1]}' }} }}
            }},
            series: [
                {{
                    name: '收入',
                    type: 'line',
                    smooth: true,
                    data: {revenue},
                    lineStyle: {{ color: '{colors['primary'][0]}', width: 3 }},
                    itemStyle: {{ color: '{colors['primary'][0]}' }},
                    areaStyle: {{ color: new echarts.graphic.LinearGradient(0, 0, 0, 1, [{{ offset: 0, color: '{colors['primary'][0]}40' }}, {{ offset: 1, color: '{colors['primary'][0]}10' }}]) }}
                }},
                {{
                    name: '支出',
                    type: 'line',
                    smooth: true,
                    data: {expenses},
                    lineStyle: {{ color: '{colors['secondary'][0]}', width: 3 }},
                    itemStyle: {{ color: '{colors['secondary'][0]}' }},
                    areaStyle: {{ color: new echarts.graphic.LinearGradient(0, 0, 0, 1, [{{ offset: 0, color: '{colors['secondary'][0]}40' }}, {{ offset: 1, color: '{colors['secondary'][0]}10' }}]) }}
                }},
                {{
                    name: '净利润',
                    type: 'line',
                    smooth: true,
                    data: {profit},
                    lineStyle: {{ color: '{colors['primary'][2]}', width: 3 }},
                    itemStyle: {{ color: '{colors['primary'][2]}' }},
                    areaStyle: {{ color: new echarts.graphic.LinearGradient(0, 0, 0, 1, [{{ offset: 0, color: '{colors['primary'][2]}40' }}, {{ offset: 1, color: '{colors['primary'][2]}10' }}]) }}
                }}
            ]
        }};
        {chart_id}_chart.setOption({chart_id}_option);
        """
    
    def _create_waterfall_chart(self, chart_id: str, result: AgentAnalysisResult, colors: Dict) -> str:
        """创建成本预测瀑布图 - 基于真实数据"""
        # 从AgentAnalysisResult中提取真实成本数据
        analysis_data = result.analysis_data
        
        # 如果没有真实数据，返回错误
        if not analysis_data or 'cost_data' not in analysis_data:
            raise Exception(f"成本分析数据缺失，无法生成图表。Agent: {result.agent_name}")
        
        cost_data = analysis_data['cost_data']
        
        # 验证必要的数据字段
        required_fields = ['categories', 'values', 'auxiliary_values']
        for field in required_fields:
            if field not in cost_data:
                raise Exception(f"成本数据字段 '{field}' 缺失，无法生成图表")
        
        categories = cost_data['categories']
        values = cost_data['values']
        auxiliary_values = cost_data['auxiliary_values']
        
        return f"""
        // 成本预测瀑布图 - 基于真实数据
        const {chart_id}_chart = echarts.init(document.getElementById('{chart_id}'));
        const {chart_id}_option = {{
            title: {{
                text: '成本预测瀑布图',
                left: 'center',
                textStyle: {{
                    color: '{colors['secondary'][0]}',
                    fontSize: 18,
                    fontWeight: 'bold'
                }}
            }},
            tooltip: {{
                trigger: 'axis',
                backgroundColor: 'rgba(255, 255, 255, 0.95)',
                borderColor: '{colors['primary'][0]}',
                textStyle: {{ color: '{colors['secondary'][0]}' }}
            }},
            grid: {{
                left: '3%',
                right: '4%',
                bottom: '3%',
                containLabel: true
            }},
            xAxis: {{
                type: 'category',
                data: {categories},
                axisLine: {{ lineStyle: {{ color: '{colors['accent'][2]}' }} }},
                axisLabel: {{ color: '{colors['secondary'][2]}', rotate: 45 }}
            }},
            yAxis: {{
                type: 'value',
                axisLine: {{ lineStyle: {{ color: '{colors['accent'][2]}' }} }},
                axisLabel: {{ color: '{colors['secondary'][2]}' }},
                splitLine: {{ lineStyle: {{ color: '{colors['accent'][1]}' }} }}
            }},
            series: [
                {{
                    name: '辅助',
                    type: 'bar',
                    stack: 'total',
                    itemStyle: {{ color: 'transparent' }},
                    data: {auxiliary_values}
                }},
                {{
                    name: '成本',
                    type: 'bar',
                    stack: 'total',
                    itemStyle: {{
                        color: function(params) {{
                            const colors = ['{colors['primary'][0]}', '{colors['primary'][1]}', '{colors['primary'][2]}', '{colors['primary'][3]}', '{colors['secondary'][0]}', '{colors['primary'][0]}'];
                            return colors[params.dataIndex % colors.length];
                        }}
                    }},
                    data: {values}
                }}
            ]
        }};
        {chart_id}_chart.setOption({chart_id}_option);
        """
    
    def _create_performance_radar_chart(self, chart_id: str, result: AgentAnalysisResult, colors: Dict) -> str:
        """创建效能评估雷达图 - 基于真实数据"""
        # 从AgentAnalysisResult中提取真实效能数据
        analysis_data = result.analysis_data
        
        # 如果没有真实数据，返回错误
        if not analysis_data or 'performance_data' not in analysis_data:
            raise Exception(f"效能评估数据缺失，无法生成图表。Agent: {result.agent_name}")
        
        performance_data = analysis_data['performance_data']
        
        # 验证必要的数据字段
        required_fields = ['indicators', 'current_values', 'target_values']
        for field in required_fields:
            if field not in performance_data:
                raise Exception(f"效能数据字段 '{field}' 缺失，无法生成图表")
        
        indicators = performance_data['indicators']
        current_values = performance_data['current_values']
        target_values = performance_data['target_values']
        
        return f"""
        // 效能评估雷达图 - 基于真实数据
        const {chart_id}_chart = echarts.init(document.getElementById('{chart_id}'));
        const {chart_id}_option = {{
            title: {{
                text: '效能评估雷达图',
                left: 'center',
                textStyle: {{
                    color: '{colors['secondary'][0]}',
                    fontSize: 18,
                    fontWeight: 'bold'
                }}
            }},
            tooltip: {{
                backgroundColor: 'rgba(255, 255, 255, 0.95)',
                borderColor: '{colors['primary'][0]}',
                textStyle: {{ color: '{colors['secondary'][0]}' }}
            }},
            radar: {{
                indicator: {indicators},
                shape: 'polygon',
                radius: '60%',
                axisName: {{
                    color: '{colors['secondary'][1]}',
                    fontSize: 12
                }},
                splitLine: {{
                    lineStyle: {{ color: '{colors['accent'][2]}' }}
                }},
                splitArea: {{
                    areaStyle: {{
                        color: ['{colors['accent'][0]}20', '{colors['accent'][1]}20']
                    }}
                }}
            }},
            series: [
                {{
                    name: '当前效能',
                    type: 'radar',
                    data: [
                        {{
                            value: {current_values},
                            name: '当前效能',
                            itemStyle: {{ color: '{colors['primary'][0]}' }},
                            areaStyle: {{ color: '{colors['primary'][0]}40' }},
                            lineStyle: {{ color: '{colors['primary'][0]}', width: 2 }}
                        }},
                        {{
                            value: {target_values},
                            name: '目标效能',
                            itemStyle: {{ color: '{colors['primary'][2]}' }},
                            areaStyle: {{ color: '{colors['primary'][2]}20' }},
                            lineStyle: {{ color: '{colors['primary'][2]}', width: 2, type: 'dashed' }}
                        }}
                    ]
                }}
            ]
        }};
        {chart_id}_chart.setOption({chart_id}_option);
        """
    
    def _create_heatmap_chart(self, chart_id: str, result: AgentAnalysisResult, colors: Dict) -> str:
        """创建运维知识分布热力图 - 基于真实数据"""
        # 从AgentAnalysisResult中提取真实热力图数据
        analysis_data = result.analysis_data
        
        # 如果没有真实数据，返回错误
        if not analysis_data or 'heatmap_data' not in analysis_data:
            raise Exception(f"热力图数据缺失，无法生成图表。Agent: {result.agent_name}")
        
        heatmap_data = analysis_data['heatmap_data']
        
        # 验证必要的数据字段
        required_fields = ['x_categories', 'y_categories', 'data_points', 'min_value', 'max_value']
        for field in required_fields:
            if field not in heatmap_data:
                raise Exception(f"热力图数据字段 '{field}' 缺失，无法生成图表")
        
        x_categories = heatmap_data['x_categories']
        y_categories = heatmap_data['y_categories']
        data_points = heatmap_data['data_points']
        min_value = heatmap_data['min_value']
        max_value = heatmap_data['max_value']
        
        return f"""
        // 运维知识分布热力图 - 基于真实数据
        const {chart_id}_chart = echarts.init(document.getElementById('{chart_id}'));
        const {chart_id}_option = {{
            title: {{
                text: '运维知识分布热力图',
                left: 'center',
                textStyle: {{
                    color: '{colors['secondary'][0]}',
                    fontSize: 18,
                    fontWeight: 'bold'
                }}
            }},
            tooltip: {{
                position: 'top',
                backgroundColor: 'rgba(255, 255, 255, 0.95)',
                borderColor: '{colors['primary'][0]}',
                textStyle: {{ color: '{colors['secondary'][0]}' }}
            }},
            grid: {{
                height: '50%',
                top: '10%'
            }},
            xAxis: {{
                type: 'category',
                data: {x_categories},
                splitArea: {{ show: true }},
                axisLabel: {{ color: '{colors['secondary'][2]}' }}
            }},
            yAxis: {{
                type: 'category',
                data: {y_categories},
                splitArea: {{ show: true }},
                axisLabel: {{ color: '{colors['secondary'][2]}' }}
            }},
            visualMap: {{
                min: {min_value},
                max: {max_value},
                calculable: true,
                orient: 'horizontal',
                left: 'center',
                bottom: '5%',
                inRange: {{
                    color: ['{colors['accent'][0]}', '{colors['primary'][3]}', '{colors['primary'][0]}']
                }},
                textStyle: {{ color: '{colors['secondary'][2]}' }}
            }},
            series: [
                {{
                    name: '知识覆盖度',
                    type: 'heatmap',
                    data: {data_points},
                    label: {{
                        show: true,
                        color: '{colors['secondary'][0]}'
                    }},
                    emphasis: {{
                        itemStyle: {{
                            shadowBlur: 10,
                            shadowColor: 'rgba(0, 0, 0, 0.5)'
                        }}
                    }}
                }}
            ]
        }};
        {chart_id}_chart.setOption({chart_id}_option);
        """
    
    def _create_scatter_chart(self, chart_id: str, result: AgentAnalysisResult, colors: Dict) -> str:
        """创建数据关联分析散点图 - 基于真实数据"""
        # 从AgentAnalysisResult中提取真实散点图数据
        analysis_data = result.analysis_data
        
        # 如果没有真实数据，返回错误
        if not analysis_data or 'scatter_data' not in analysis_data:
            raise Exception(f"散点图数据缺失，无法生成图表。Agent: {result.agent_name}")
        
        scatter_data = analysis_data['scatter_data']
        
        # 验证必要的数据字段
        required_fields = ['data_points', 'x_axis_name', 'y_axis_name', 'series_name']
        for field in required_fields:
            if field not in scatter_data:
                raise Exception(f"散点图数据字段 '{field}' 缺失，无法生成图表")
        
        data_points = scatter_data['data_points']
        x_axis_name = scatter_data['x_axis_name']
        y_axis_name = scatter_data['y_axis_name']
        series_name = scatter_data['series_name']
        
        return f"""
        // 数据关联分析散点图 - 基于真实数据
        const {chart_id}_chart = echarts.init(document.getElementById('{chart_id}'));
        const {chart_id}_option = {{
            title: {{
                text: '数据关联分析',
                left: 'center',
                textStyle: {{
                    color: '{colors['secondary'][0]}',
                    fontSize: 18,
                    fontWeight: 'bold'
                }}
            }},
            tooltip: {{
                trigger: 'item',
                backgroundColor: 'rgba(255, 255, 255, 0.95)',
                borderColor: '{colors['primary'][0]}',
                textStyle: {{ color: '{colors['secondary'][0]}' }}
            }},
            grid: {{
                left: '3%',
                right: '7%',
                bottom: '3%',
                containLabel: true
            }},
            xAxis: {{
                type: 'value',
                name: '{x_axis_name}',
                axisLine: {{ lineStyle: {{ color: '{colors['accent'][2]}' }} }},
                axisLabel: {{ color: '{colors['secondary'][2]}' }},
                splitLine: {{ lineStyle: {{ color: '{colors['accent'][1]}' }} }}
            }},
            yAxis: {{
                type: 'value',
                name: '{y_axis_name}',
                axisLine: {{ lineStyle: {{ color: '{colors['accent'][2]}' }} }},
                axisLabel: {{ color: '{colors['secondary'][2]}' }},
                splitLine: {{ lineStyle: {{ color: '{colors['accent'][1]}' }} }}
            }},
            series: [
                {{
                    name: '{series_name}',
                    type: 'scatter',
                    data: {data_points},
                    symbolSize: function(data) {{ return Math.sqrt(data[2] || 100) / 2; }},
                    itemStyle: {{
                        color: '{colors['primary'][0]}',
                        opacity: 0.8
                    }},
                    emphasis: {{
                        itemStyle: {{
                            color: '{colors['primary'][2]}',
                            borderColor: '{colors['secondary'][0]}',
                            borderWidth: 2
                        }}
                    }}
                }}
            ]
        }};
        {chart_id}_chart.setOption({chart_id}_option);
        """
    
    def _create_staff_bar_chart(self, chart_id: str, result: AgentAnalysisResult, colors: Dict) -> str:
        """创建人员效能对比柱状图 - 基于真实数据"""
        # 从AgentAnalysisResult中提取真实柱状图数据
        analysis_data = result.analysis_data
        
        # 如果没有真实数据，返回错误
        if not analysis_data or 'bar_chart_data' not in analysis_data:
            raise Exception(f"柱状图数据缺失，无法生成图表。Agent: {result.agent_name}")
        
        bar_data = analysis_data['bar_chart_data']
        
        # 验证必要的数据字段
        required_fields = ['categories', 'current_values', 'target_values', 'current_name', 'target_name']
        for field in required_fields:
            if field not in bar_data:
                raise Exception(f"柱状图数据字段 '{field}' 缺失，无法生成图表")
        
        categories = bar_data['categories']
        current_values = bar_data['current_values']
        target_values = bar_data['target_values']
        current_name = bar_data['current_name']
        target_name = bar_data['target_name']
        
        return f"""
        // 人员效能对比柱状图 - 基于真实数据
        const {chart_id}_chart = echarts.init(document.getElementById('{chart_id}'));
        const {chart_id}_option = {{
            title: {{
                text: '人员效能对比',
                left: 'center',
                textStyle: {{
                    color: '{colors['secondary'][0]}',
                    fontSize: 18,
                    fontWeight: 'bold'
                }}
            }},
            tooltip: {{
                trigger: 'axis',
                backgroundColor: 'rgba(255, 255, 255, 0.95)',
                borderColor: '{colors['primary'][0]}',
                textStyle: {{ color: '{colors['secondary'][0]}' }}
            }},
            legend: {{
                data: ['{current_name}', '{target_name}'],
                top: '10%',
                textStyle: {{ color: '{colors['secondary'][1]}' }}
            }},
            grid: {{
                left: '3%',
                right: '4%',
                bottom: '3%',
                containLabel: true
            }},
            xAxis: {{
                type: 'category',
                data: {categories},
                axisLine: {{ lineStyle: {{ color: '{colors['accent'][2]}' }} }},
                axisLabel: {{ color: '{colors['secondary'][2]}', rotate: 45 }}
            }},
            yAxis: {{
                type: 'value',
                axisLine: {{ lineStyle: {{ color: '{colors['accent'][2]}' }} }},
                axisLabel: {{ color: '{colors['secondary'][2]}' }},
                splitLine: {{ lineStyle: {{ color: '{colors['accent'][1]}' }} }}
            }},
            series: [
                {{
                    name: '{current_name}',
                    type: 'bar',
                    data: {current_values},
                    itemStyle: {{
                        color: new echarts.graphic.LinearGradient(0, 0, 0, 1, [
                            {{ offset: 0, color: '{colors['primary'][0]}' }},
                            {{ offset: 1, color: '{colors['primary'][1]}' }}
                        ])
                    }}
                }},
                {{
                    name: '{target_name}',
                    type: 'bar',
                    data: {target_values},
                    itemStyle: {{
                        color: new echarts.graphic.LinearGradient(0, 0, 0, 1, [
                            {{ offset: 0, color: '{colors['primary'][2]}40' }},
                            {{ offset: 1, color: '{colors['primary'][3]}40' }}
                        ]),
                        borderColor: '{colors['primary'][2]}',
                        borderWidth: 2,
                        borderType: 'dashed'
                    }}
                }}
            ]
        }};
        {chart_id}_chart.setOption({chart_id}_option);
        """
    
    def _create_modern_pie_chart(self, chart_id: str, result: AgentAnalysisResult, colors: Dict) -> str:
        """创建现代饼图 - 基于真实数据"""
        # 从AgentAnalysisResult中提取真实饼图数据
        analysis_data = result.analysis_data
        
        # 如果没有真实数据，返回错误
        if not analysis_data or 'pie_chart_data' not in analysis_data:
            raise Exception(f"饼图数据缺失，无法生成图表。Agent: {result.agent_name}")
        
        pie_data = analysis_data['pie_chart_data']
        
        # 验证必要的数据字段
        required_fields = ['data_items', 'series_name']
        for field in required_fields:
            if field not in pie_data:
                raise Exception(f"饼图数据字段 '{field}' 缺失，无法生成图表")
        
        data_items = pie_data['data_items']
        series_name = pie_data['series_name']
        
        # 为数据项添加颜色
        color_palette = [colors['primary'][0], colors['primary'][1], colors['primary'][2], colors['secondary'][0]]
        for i, item in enumerate(data_items):
            if 'itemStyle' not in item:
                item['itemStyle'] = {'color': color_palette[i % len(color_palette)]}
        
        return f"""
        // 现代饼图 - 基于真实数据
        const {chart_id}_chart = echarts.init(document.getElementById('{chart_id}'));
        const {chart_id}_option = {{
            title: {{
                text: '{result.agent_name}分析结果',
                left: 'center',
                textStyle: {{
                    color: '{colors['secondary'][0]}',
                    fontSize: 18,
                    fontWeight: 'bold'
                }}
            }},
            tooltip: {{
                trigger: 'item',
                backgroundColor: 'rgba(255, 255, 255, 0.95)',
                borderColor: '{colors['primary'][0]}',
                textStyle: {{ color: '{colors['secondary'][0]}' }}
            }},
            legend: {{
                orient: 'vertical',
                left: 'left',
                textStyle: {{ color: '{colors['secondary'][1]}' }}
            }},
            series: [
                {{
                    name: '{series_name}',
                    type: 'pie',
                    radius: ['40%', '70%'],
                    center: ['60%', '50%'],
                    avoidLabelOverlap: false,
                    itemStyle: {{
                        borderRadius: 10,
                        borderColor: '#fff',
                        borderWidth: 2
                    }},
                    label: {{
                        show: false,
                        position: 'center'
                    }},
                    emphasis: {{
                        label: {{
                            show: true,
                            fontSize: '16',
                            fontWeight: 'bold',
                            color: '{colors['secondary'][0]}'
                        }}
                    }},
                    labelLine: {{
                        show: false
                    }},
                    data: {data_items}
                }}
            ]
        }};
        {chart_id}_chart.setOption({chart_id}_option);
        """
    
    def _format_recommendations(self, recommendations: List[str]) -> str:
        """格式化建议列表"""
        if not recommendations:
            return "<p>暂无具体建议</p>"
        
        formatted = "<ul>"
        for rec in recommendations:
            formatted += f"<li>{rec}</li>"
        formatted += "</ul>"
        
        return formatted
    
    def _format_risk_assessment(self, risk_data: Dict[str, Any]) -> str:
        """格式化风险评估"""
        if not risk_data:
            return "<p>未识别到重大风险</p>"
        
        formatted = "<ul>"
        for risk, description in risk_data.items():
            formatted += f"<li><strong>{risk}</strong>: {description}</li>"
        formatted += "</ul>"
        
        return formatted
    
    def _format_confidence_metrics(self, overall_confidence: float) -> str:
        """格式化置信度指标"""
        confidence_percentage = overall_confidence * 100
        
        return f"""
        <div class="confidence-meter">
            <span>综合分析置信度:</span>
            <div class="confidence-bar">
                <div class="confidence-fill" style="width: {confidence_percentage:.0f}%"></div>
            </div>
            <span>{confidence_percentage:.0f}%</span>
        </div>
        <p>置信度说明：基于各专业Agent分析结果的加权平均值，反映综合分析的可靠程度。</p>
        """

# ================================
# 4. 工厂函数
# ================================

def create_report_generator_agent() -> ReportGeneratorAgent:
    """创建报告生成专家Agent实例"""
    return ReportGeneratorAgent()

# ================================
# 5. 模块导出
# ================================

__all__ = [
    'ReportGeneratorAgent',
    'AgentAnalysisResult', 
    'ReportSection',
    'ComprehensiveReport',
    'create_report_generator_agent'
]